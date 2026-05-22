from __future__ import annotations

import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from bs4 import BeautifulSoup


def _pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def _has_math_markers(text: str) -> bool:
    markers = [
        "$",
        "\\frac",
        "\\sqrt",
        "\\sum",
        "\\int",
        "\\log",
        "\\Delta",
        "\\alpha",
        "\\beta",
        "\\sigma",
        "\\mu",
        "^",
        "_",
    ]
    return any(marker in text for marker in markers)


def _count_omml_equations(docx_path: str) -> int:
    """
    Conta equações nativas do Word em OMML.

    Isso é útil para saber se o DOCX tem equações editáveis.
    """
    path = Path(docx_path)

    if path.suffix.lower() != ".docx":
        return 0

    count = 0

    try:
        with zipfile.ZipFile(path, "r") as z:
            xml_files = [
                name for name in z.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]

            for xml_file in xml_files:
                xml_content = z.read(xml_file).decode("utf-8", errors="ignore")
                soup = BeautifulSoup(xml_content, "xml")

                count += len(soup.find_all("m:oMath"))
                count += len(soup.find_all("m:oMathPara"))

    except Exception:
        return 0

    return count


def _load_docx_with_pandoc(docx_path: str) -> str:
    """
    Usa Pandoc para converter DOCX em Markdown.

    Quando as equações do Word são editáveis, o Pandoc geralmente
    preserva matemática em formato LaTeX.
    """
    path = Path(docx_path)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        output_md = tmpdir / "output.md"
        media_dir = tmpdir / "media"

        cmd = [
            "pandoc",
            str(path),
            "-t",
            "markdown+tex_math_dollars",
            "--wrap=none",
            f"--extract-media={media_dir}",
            "-o",
            str(output_md),
        ]

        subprocess.run(
            cmd,
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )

        if not output_md.exists():
            return ""

        text = output_md.read_text(encoding="utf-8", errors="ignore")

        # Marca imagens extraídas. Algumas podem conter fórmulas/gráficos.
        images = []
        if media_dir.exists():
            for img in media_dir.rglob("*"):
                if img.suffix.lower() in [".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"]:
                    images.append(img.name)

        if images:
            text += "\n\n[IMAGENS EXTRAÍDAS DO DOCUMENTO]\n"
            for img in images:
                text += f"\n[imagem detectada: {img} - possível fórmula, tabela, gráfico ou figura]"

        return text.strip()


def _load_docx_with_python_docx(docx_path: str) -> str:
    """
    Fallback simples com python-docx.

    Atenção: esse método normalmente perde equações do Word.
    Ele é usado apenas se Pandoc não estiver disponível ou falhar.
    """
    doc = Document(docx_path)

    parts: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Tabelas do Word
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts).strip()


def load_docx_text(docx_path: str) -> str:
    """
    Carrega DOCX priorizando preservação matemática.

    Ordem:
    1. Tenta Pandoc, que pode preservar equações como LaTeX.
    2. Se falhar, usa python-docx.
    3. Se houver OMML detectado mas não convertido, adiciona aviso.
    """

    omml_count = _count_omml_equations(docx_path)

    text = ""

    if _pandoc_available():
        try:
            text = _load_docx_with_pandoc(docx_path)
        except Exception as e:
            text = ""
            pandoc_error = str(e)
        else:
            pandoc_error = ""
    else:
        pandoc_error = "Pandoc não encontrado no sistema."

    if not text:
        text = _load_docx_with_python_docx(docx_path)

    # Se havia equações OMML e o texto final não tem sinais de LaTeX,
    # provavelmente as equações foram perdidas.
    if omml_count > 0 and not _has_math_markers(text):
        text += (
            "\n\n[AVISO: foram detectadas equações editáveis do Word, "
            "mas elas não foram convertidas para LaTeX. "
            "Verifique se o Pandoc está instalado corretamente.]"
        )

    if pandoc_error:
        text += f"\n\n[AVISO PANDOC: {pandoc_error}]"

    return text.strip()